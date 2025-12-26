#!/usr/bin/env python3
"""Convert Kalman Optimization Report from Markdown to DOCX."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

# Paths
REPORT_MD = Path(__file__).parent.parent / "docs/reports/KALMAN_HYBRID_OPTIMIZATION_REPORT.md"
REPORT_DOCX = Path(__file__).parent.parent / "docs/reports/KALMAN_HYBRID_OPTIMIZATION_REPORT.docx"


def create_styles(doc):
    """Create custom styles for the document."""
    styles = doc.styles

    # Title style
    if 'Report Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

    # Subtitle style
    if 'Report Subtitle' not in [s.name for s in styles]:
        subtitle_style = styles.add_style('Report Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.italic = True
        subtitle_style.font.color.rgb = RGBColor(80, 80, 80)
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(24)


def parse_markdown(content: str) -> list:
    """Parse markdown content into structured elements."""
    elements = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
        elif line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))

        # Horizontal rule
        elif line.strip() == '---':
            elements.append(('hr', ''))

        # Code blocks
        elif line.strip().startswith('```'):
            code_lines = []
            lang = line.strip()[3:]
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append(('code', '\n'.join(code_lines)))

        # Tables
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                if '---' not in lines[i]:
                    table_lines.append(lines[i])
                i += 1
            elements.append(('table', table_lines))
            continue

        # Bold/metadata lines
        elif line.startswith('**') and line.endswith('**'):
            elements.append(('bold', line.strip('*')))

        # List items
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            elements.append(('list', line.strip()[2:]))
        elif re.match(r'^\d+\. ', line.strip()):
            elements.append(('numlist', re.sub(r'^\d+\. ', '', line.strip())))

        # Regular paragraph
        else:
            # Collect multi-line paragraph
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('|') and not lines[i].startswith('```') and not lines[i].startswith('-') and not lines[i].strip() == '---':
                para_lines.append(lines[i])
                i += 1
            elements.append(('para', ' '.join(para_lines)))
            continue

        i += 1

    return elements


def add_table(doc, table_lines):
    """Add a table to the document."""
    # Parse table
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()  # Space after table


def convert_to_docx(md_path: Path, docx_path: Path):
    """Convert markdown file to DOCX."""
    # Read markdown
    content = md_path.read_text()

    # Create document
    doc = Document()
    create_styles(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Parse and add content
    elements = parse_markdown(content)

    for elem_type, elem_content in elements:
        if elem_type == 'h1':
            p = doc.add_heading(elem_content, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif elem_type == 'h2':
            doc.add_heading(elem_content, level=1)
        elif elem_type == 'h3':
            doc.add_heading(elem_content, level=2)
        elif elem_type == 'h4':
            doc.add_heading(elem_content, level=3)
        elif elem_type == 'hr':
            doc.add_paragraph('_' * 80)
        elif elem_type == 'code':
            p = doc.add_paragraph()
            run = p.add_run(elem_content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
        elif elem_type == 'table':
            add_table(doc, elem_content)
        elif elem_type == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(elem_content)
            run.bold = True
        elif elem_type == 'list':
            doc.add_paragraph(elem_content, style='List Bullet')
        elif elem_type == 'numlist':
            doc.add_paragraph(elem_content, style='List Number')
        elif elem_type == 'para':
            # Handle inline formatting
            p = doc.add_paragraph()
            # Simple inline bold/code handling
            text = elem_content
            # Replace **text** with bold
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for idx, part in enumerate(parts):
                if idx % 2 == 0:
                    # Handle inline code
                    code_parts = re.split(r'`(.+?)`', part)
                    for cidx, cpart in enumerate(code_parts):
                        run = p.add_run(cpart)
                        if cidx % 2 == 1:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                else:
                    run = p.add_run(part)
                    run.bold = True

    # Save document
    doc.save(docx_path)
    print(f"Report saved to: {docx_path}")
    print(f"Word count: ~{len(content.split())} words")


if __name__ == "__main__":
    convert_to_docx(REPORT_MD, REPORT_DOCX)
